import io
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from rest_framework.test import APITestCase
from rest_framework import status
from .models import Resume

User = get_user_model()


class ResumeUploadTests(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="testuser",
            email="testuser@example.com",
            password="TestPassword123!",
            first_name="Ritesh",
            last_name="Pandey",
        )
        self.client.force_authenticate(user=self.user)

    def test_upload_resume_docx(self):
        # Create a sample docx in memory
        import docx
        doc = docx.Document()
        doc.add_heading("Ritesh Pandey", 0)
        doc.add_paragraph("Senior Full-Stack Engineer | ritesh.pandey@example.com | +91 98765 43210")
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph("Experienced software engineer specializing in scalable cloud architectures and React.")
        doc.add_heading("Experience", level=1)
        doc.add_paragraph("Lead Engineer - Swiggy (2021-Present)")
        doc.add_paragraph("• Architected order dispatch microservices handling 120k req/sec.")
        doc.add_paragraph("• Cut infrastructure cost by 34% through Kubernetes autoscaling.")
        doc.add_heading("Education", level=1)
        doc.add_paragraph("B.Tech in Computer Science - IIT Madras (2017-2021)")
        doc.add_heading("Skills", level=1)
        doc.add_paragraph("React, TypeScript, Node.js, Python, PostgreSQL, AWS, Docker")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        uploaded = SimpleUploadedFile(
            "ritesh_resume.docx",
            file_stream.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response = self.client.post("/api/resumes/upload/", {"file": uploaded}, format="multipart")
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        data = response.json()

        self.assertEqual(data["source"], "upload")
        self.assertEqual(data["status"], "ready")
        self.assertIn("experiences", data)
        self.assertIn("education", data)
        self.assertIn("skills", data)
        self.assertIn("raw_ai_extraction", data)

        # Check DB
        resume = Resume.objects.get(id=data["id"])
        self.assertEqual(resume.user, self.user)
        self.assertEqual(resume.experiences.count(), len(data["experiences"]))

    def test_upload_unsupported_file_type(self):
        uploaded = SimpleUploadedFile(
            "bad_file.txt",
            b"Plain text content",
            content_type="text/plain",
        )
        response = self.client.post("/api/resumes/upload/", {"file": uploaded}, format="multipart")
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)

    def test_upload_unauthenticated(self):
        self.client.force_authenticate(user=None)
        uploaded = SimpleUploadedFile(
            "test.pdf",
            b"%PDF-1.4 dummy content",
            content_type="application/pdf",
        )
        response = self.client.post("/api/resumes/upload/", {"file": uploaded}, format="multipart")
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)

